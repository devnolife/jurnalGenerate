from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(document, font_name, font_size):
    style = document.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    r_fonts = style.element.rPr.rFonts
    r_fonts.set(qn('w:eastAsia'), font_name)

def add_section(document, title, content):
    heading = document.add_heading(level=2)
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(14)

    for paragraph in content:
        document.add_paragraph(paragraph)

sections_content = {
    "Pendahuluan": [
        "Media sosial telah menjadi bagian integral dari kehidupan sehari-hari, terutama di kalangan mahasiswa. Penggunaan media sosial yang intensif dapat mempengaruhi berbagai aspek kehidupan, termasuk kinerja akademik. Penelitian ini bertujuan untuk memahami dampak penggunaan media sosial terhadap kinerja akademik mahasiswa di Universitas X.",
        "Rumusan Masalah: 1. Bagaimana tingkat penggunaan media sosial di kalangan mahasiswa Universitas X? 2. Apakah ada hubungan antara penggunaan media sosial dan kinerja akademik mahasiswa? 3. Faktor-faktor apa saja yang mempengaruhi hubungan antara penggunaan media sosial dan kinerja akademik?",
        "Tujuan Penelitian: 1. Mengukur tingkat penggunaan media sosial di kalangan mahasiswa. 2. Menganalisis hubungan antara penggunaan media sosial dan kinerja akademik. 3. Mengidentifikasi faktor-faktor yang mempengaruhi hubungan tersebut.",
        "Manfaat Penelitian: Hasil penelitian ini diharapkan dapat memberikan wawasan bagi mahasiswa dan pendidik tentang pengaruh media sosial terhadap kinerja akademik serta memberikan rekomendasi untuk penggunaan media sosial yang lebih bijak."
    ],
    "Tinjauan Pustaka": [
        "Penelitian tentang dampak media sosial terhadap kinerja akademik telah dilakukan oleh berbagai peneliti. Menurut Smith (2019), penggunaan media sosial yang berlebihan dapat menyebabkan gangguan konsentrasi yang berdampak negatif pada kinerja akademik. Di sisi lain, penelitian oleh Johnson (2020) menunjukkan bahwa media sosial dapat digunakan sebagai alat pembelajaran yang efektif jika digunakan dengan bijak. Teori yang mendasari penelitian ini adalah teori penggunaan dan gratifikasi, yang menyatakan bahwa individu menggunakan media untuk memenuhi kebutuhan dan mendapatkan gratifikasi tertentu."
    ],
    "Metodologi": [
        "Penelitian ini menggunakan desain penelitian kuantitatif dengan metode survei. Populasi penelitian adalah mahasiswa Universitas X yang aktif menggunakan media sosial. Sampel diambil secara acak sebanyak 200 responden. Teknik pengumpulan data menggunakan kuesioner yang terdiri dari pertanyaan tertutup dan terbuka. Data yang dikumpulkan dianalisis menggunakan teknik statistik deskriptif dan inferensial untuk menguji hipotesis penelitian."
    ],
    "Hasil dan Pembahasan": [
        "Hasil penelitian menunjukkan bahwa 70% mahasiswa menggunakan media sosial lebih dari 3 jam per hari. Analisis korelasi menunjukkan adanya hubungan negatif yang signifikan antara penggunaan media sosial dan kinerja akademik (r = -0.45, p < 0.05). Faktor-faktor seperti manajemen waktu dan motivasi belajar ditemukan memoderasi hubungan tersebut. Mahasiswa yang memiliki keterampilan manajemen waktu yang baik cenderung tidak terpengaruh oleh penggunaan media sosial.",
        "Temuan ini sejalan dengan penelitian Smith (2019) yang menyatakan bahwa penggunaan media sosial dapat mengganggu konsentrasi belajar. Namun, hasil ini juga menunjukkan pentingnya manajemen waktu dan motivasi belajar sebagai faktor yang dapat mengurangi dampak negatif media sosial. Oleh karena itu, intervensi yang berfokus pada peningkatan keterampilan manajemen waktu dan motivasi belajar dapat membantu mahasiswa menggunakan media sosial dengan lebih bijak."
    ],
    "Kesimpulan": [
        "Penelitian ini menyimpulkan bahwa penggunaan media sosial yang berlebihan dapat berdampak negatif pada kinerja akademik mahasiswa Universitas X. Namun, faktor-faktor seperti manajemen waktu dan motivasi belajar dapat memoderasi dampak tersebut. Disarankan agar mahasiswa meningkatkan keterampilan manajemen waktu dan motivasi belajar untuk meminimalkan dampak negatif media sosial. Penelitian lebih lanjut diperlukan untuk mengeksplorasi intervensi yang efektif dalam konteks ini."
    ]
}

doc = Document()
set_font(doc, 'Times New Roman', 12)
for section_title, section_content in sections_content.items():
    add_section(doc, section_title, section_content)
doc.save("journal_draft.docx")

print("Journal draft generated!")
