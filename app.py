import PyPDF2
import nltk
from docx import Document
from langdetect import detect
import re

nltk.download('punkt')

def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def identify_sections(text):
    sections = {}
    section_patterns = {
        "Pendahuluan": re.compile(r"\bPendahuluan\b", re.IGNORECASE),
        "Tinjauan Pustaka": re.compile(r"\bTinjauan Pustaka\b", re.IGNORECASE),
        "Metodologi": re.compile(r"\bMetodologi\b", re.IGNORECASE),
        "Hasil dan Pembahasan": re.compile(r"\bHasil dan Pembahasan\b", re.IGNORECASE),
        "Kesimpulan": re.compile(r"\bKesimpulan\b", re.IGNORECASE)
    }

    sentences = nltk.sent_tokenize(text)
    current_section = None
    for sentence in sentences:
        for section, pattern in section_patterns.items():
            if pattern.search(sentence):
                current_section = section
                sections[current_section] = []
        if current_section:
            sections[current_section].append(sentence)

    return sections

thesis_pdf = "thesis.pdf"
journal_template = "journal_template.docx"
thesis_text = extract_text_from_pdf(thesis_pdf)
language = detect(thesis_text)
if language != 'id':
    print("The document is not in Indonesian.")
else:
    print("Detected Indonesian language.")
    sections = identify_sections(thesis_text)
    doc = Document(journal_template)
    for section_name, content in sections.items():
        paragraph = doc.add_paragraph(section_name, style='Heading 2')
        for sentence in content:
            paragraph.add_run(sentence).add_break()
    doc.save("journal_draft.docx")
    print("Journal draft generated!")
